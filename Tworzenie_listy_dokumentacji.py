import os
import sys
from pathlib import Path
from tkinter import filedialog, Tk
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL


# Folder aplikacji (dla PyInstaller)
app_folder = Path(getattr(sys, '_MEIPASS', os.path.abspath(os.path.dirname(__file__))))

# Wybór folderu źródłowego
Tk().withdraw()
folder_path = filedialog.askdirectory(title="Wybierz folder główny")
if not folder_path:
    print("Nie wybrano folderu.")
    exit()

# Tworzenie dokumentu Word
doc = Document()

# Wąskie marginesy
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Logo i dane firmy w tabeli
logo_path = app_folder / "Logo.png"
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.allow_autofit = False
table.columns[0].width = Inches(2.2)
table.columns[1].width = Inches(4.8)

# Usunięcie ramek z tabeli
tbl = table._tbl
tblPr = tbl.tblPr  # Uzyskanie dostępu do właściwości tabeli
tblBorders = OxmlElement('w:tblBorders')
for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
    border = OxmlElement(f"w:{border_name}")
    border.set(qn("w:val"), "nil")
    tblBorders.append(border)
tblPr.append(tblBorders)


# Komórka z logo
cell_logo = table.cell(0, 0)
cell_logo.vertical_alignment = WD_ALIGN_PARAGRAPH.LEFT
paragraph_logo = cell_logo.paragraphs[0]
paragraph_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
if logo_path.exists():
    run_logo = paragraph_logo.add_run()
    run_logo.add_picture(str(logo_path), width=Inches(1.2))

# Komórka z danymi firmy
header_info = [
    "MP ENERGY Sp. z o.o.",
    "ul. Warszawska 43, piętro 4",
    "61-028 Poznań",
    "biuro@mpenergy.pl",
    "https://mpenergy.pl/"
]

cell_text = table.cell(0, 1)
cell_text.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
paragraph_text = cell_text.paragraphs[0]
for i, line in enumerate(header_info):
    run = paragraph_text.add_run(line)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    if "@" in line or line.startswith("http"):
        run.font.underline = True
    if i < len(header_info) - 1:
        paragraph_text.add_run("\n")
paragraph_text.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Dodanie odstępu po nagłówku
doc.add_paragraph("Lista dokumentów").runs[0].bold = True
doc.paragraphs[-1].runs[0].font.size = Pt(16)

# Rekurencyjne przeszukiwanie folderów i tworzenie listy
def process_folder(path, numbering_stack=None, level=0):
    if numbering_stack is None:
        numbering_stack = []

    folders = []
    files = []

    for item in sorted(os.listdir(path)):
        item_path = os.path.join(path, item)
        if os.path.isdir(item_path):
            folders.append(item)
        elif os.path.isfile(item_path):
            files.append(item)

    for index, folder_name in enumerate(folders, 1):
        current_numbering = numbering_stack + [index]
        numbering_str = '.'.join(str(n) for n in current_numbering)

        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Pt(level * 20)
        para.add_run(f"{numbering_str}. {folder_name}").bold = True

        # Rekurencja
        folder_path = os.path.join(path, folder_name)
        process_folder(folder_path, current_numbering, level + 1)

    for i, file_name in enumerate(files, 1):
        current_numbering = numbering_stack + [len(folders) + i]
        numbering_str = '.'.join(str(n) for n in current_numbering)

        name_without_ext = os.path.splitext(file_name)[0]
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Pt(level * 20 + 20)
        run = para.add_run(f"{numbering_str}. {name_without_ext}")
        run.font.color.rgb = RGBColor(0, 0, 0)

        if file_name.lower().endswith(".txt"):  

            missing_run = para.add_run(" (No document)")
            missing_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

#  Start
process_folder(folder_path)

#  Zapis na pulpit
desktop = Path.home() / "Desktop"
output_path = desktop / "Lista dokumentacji.docx"
doc.save(output_path)

print(f"✔ Zapisano dokument: {output_path}")
